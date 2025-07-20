"""Text Processing Module - Handles multiple input formats and context management."""

import io
import csv
import json
import pandas as pd
from pathlib import Path
from typing import Dict, Any, List, Optional, Union, Tuple
from dataclasses import dataclass
import tiktoken

try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


@dataclass
class ProcessedText:
    content: str
    metadata: Dict[str, Any]
    chunks: List[str]
    total_tokens: int
    format_type: str
    confidence: float


class TextProcessor:
    """Processes various text formats and prepares them for LLM consumption."""
    
    def __init__(self, max_tokens_per_chunk: int = 128000):
        """
        Initialize text processor.
        
        Args:
            max_tokens_per_chunk: Maximum tokens per chunk for processing
        """
        self.max_tokens_per_chunk = max_tokens_per_chunk
        
        # Initialize tokenizer for token counting
        try:
            self.tokenizer = tiktoken.get_encoding("cl100k_base")
        except Exception:
            self.tokenizer = None
    
    def process_input(self, 
                     input_data: Union[str, bytes, io.BytesIO], 
                     file_type: Optional[str] = None,
                     filename: Optional[str] = None) -> ProcessedText:
        """
        Process input data of various formats.
        
        Args:
            input_data: Raw input data (string, bytes, or file-like object)
            file_type: Explicit file type ('pdf', 'docx', 'csv', 'txt', etc.)
            filename: Original filename for type detection
            
        Returns:
            ProcessedText object with processed content and metadata
        """
        
        # Detect file type if not provided
        if file_type is None and filename:
            file_type = self._detect_file_type(filename)
        elif file_type is None:
            file_type = 'txt'  # Default to text
        
        # Process based on file type
        if file_type.lower() in ['pdf']:
            if not PDF_AVAILABLE:
                raise ImportError("PyPDF2 is required for PDF processing")
            return self._process_pdf(input_data)
            
        elif file_type.lower() in ['docx', 'doc']:
            if not DOCX_AVAILABLE:
                raise ImportError("python-docx is required for DOCX processing")
            return self._process_docx(input_data)
            
        elif file_type.lower() in ['csv']:
            return self._process_csv(input_data)
            
        elif file_type.lower() in ['xlsx', 'xls']:
            return self._process_excel(input_data)
            
        elif file_type.lower() in ['json']:
            return self._process_json(input_data)
            
        else:
            return self._process_text(input_data)
    
    def _detect_file_type(self, filename: str) -> str:
        """Detect file type from filename extension."""
        return Path(filename).suffix.lower().lstrip('.')
    
    def _process_pdf(self, input_data: Union[bytes, io.BytesIO]) -> ProcessedText:
        """Process PDF files."""
        if isinstance(input_data, bytes):
            pdf_file = io.BytesIO(input_data)
        else:
            pdf_file = input_data
        
        try:
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            text_content = ""
            
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    text_content += f"\n--- Page {page_num + 1} ---\n{page_text}"
                except Exception as e:
                    text_content += f"\n--- Page {page_num + 1} (Error: {str(e)}) ---\n"
            
            metadata = {
                "total_pages": len(pdf_reader.pages),
                "format": "pdf",
                "extraction_method": "PyPDF2"
            }
            
            confidence = 0.8  # PDF extraction can be imperfect
            
        except Exception as e:
            text_content = f"Error processing PDF: {str(e)}"
            metadata = {"error": str(e), "format": "pdf"}
            confidence = 0.0
        
        return self._finalize_processing(text_content, metadata, "pdf", confidence)
    
    def _process_docx(self, input_data: Union[bytes, io.BytesIO]) -> ProcessedText:
        """Process DOCX files."""
        if isinstance(input_data, bytes):
            docx_file = io.BytesIO(input_data)
        else:
            docx_file = input_data
        
        try:
            doc = Document(docx_file)
            text_content = ""
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_content += para.text + "\n"
            
            # Extract tables
            for table in doc.tables:
                text_content += "\n--- Table ---\n"
                for row in table.rows:
                    row_text = " | ".join([cell.text for cell in row.cells])
                    text_content += row_text + "\n"
            
            metadata = {
                "paragraphs": len(doc.paragraphs),
                "tables": len(doc.tables),
                "format": "docx"
            }
            
            confidence = 0.9  # DOCX extraction is usually reliable
            
        except Exception as e:
            text_content = f"Error processing DOCX: {str(e)}"
            metadata = {"error": str(e), "format": "docx"}
            confidence = 0.0
        
        return self._finalize_processing(text_content, metadata, "docx", confidence)
    
    def _process_csv(self, input_data: Union[str, bytes, io.BytesIO]) -> ProcessedText:
        """Process CSV files."""
        try:
            if isinstance(input_data, bytes):
                csv_text = input_data.decode('utf-8', errors='ignore')
            elif isinstance(input_data, io.BytesIO):
                csv_text = input_data.read().decode('utf-8', errors='ignore')
            else:
                csv_text = input_data
            
            # Parse CSV
            csv_reader = csv.reader(io.StringIO(csv_text))
            rows = list(csv_reader)
            
            if not rows:
                text_content = "Empty CSV file"
                metadata = {"rows": 0, "columns": 0, "format": "csv"}
                confidence = 0.5
            else:
                # Format as structured text
                headers = rows[0] if rows else []
                text_content = f"CSV Data with {len(headers)} columns and {len(rows)} rows:\n\n"
                
                # Add headers
                text_content += "Headers: " + " | ".join(headers) + "\n\n"
                
                # Add sample rows (limit to avoid huge context)
                sample_size = min(100, len(rows) - 1)  # Skip header row
                text_content += f"Data (showing first {sample_size} rows):\n"
                
                for i, row in enumerate(rows[1:sample_size + 1], 1):
                    row_text = " | ".join([str(cell) for cell in row])
                    text_content += f"Row {i}: {row_text}\n"
                
                if len(rows) > sample_size + 1:
                    text_content += f"\n... and {len(rows) - sample_size - 1} more rows"
                
                metadata = {
                    "rows": len(rows),
                    "columns": len(headers),
                    "headers": headers,
                    "sample_size": sample_size,
                    "format": "csv"
                }
                confidence = 0.95
                
        except Exception as e:
            text_content = f"Error processing CSV: {str(e)}"
            metadata = {"error": str(e), "format": "csv"}
            confidence = 0.0
        
        return self._finalize_processing(text_content, metadata, "csv", confidence)
    
    def _process_excel(self, input_data: Union[bytes, io.BytesIO]) -> ProcessedText:
        """Process Excel files."""
        try:
            if isinstance(input_data, bytes):
                excel_file = io.BytesIO(input_data)
            else:
                excel_file = input_data
            
            # Read Excel file
            df_dict = pd.read_excel(excel_file, sheet_name=None, nrows=1000)  # Limit rows
            
            text_content = f"Excel file with {len(df_dict)} sheets:\n\n"
            
            metadata = {
                "sheets": list(df_dict.keys()),
                "total_sheets": len(df_dict),
                "format": "excel"
            }
            
            for sheet_name, df in df_dict.items():
                text_content += f"--- Sheet: {sheet_name} ---\n"
                text_content += f"Shape: {df.shape[0]} rows × {df.shape[1]} columns\n"
                text_content += f"Columns: {', '.join(df.columns.astype(str))}\n\n"
                
                # Add sample data
                sample_df = df.head(20)  # First 20 rows
                text_content += sample_df.to_string(index=False) + "\n\n"
                
                if df.shape[0] > 20:
                    text_content += f"... and {df.shape[0] - 20} more rows\n\n"
            
            confidence = 0.9
            
        except Exception as e:
            text_content = f"Error processing Excel: {str(e)}"
            metadata = {"error": str(e), "format": "excel"}
            confidence = 0.0
        
        return self._finalize_processing(text_content, metadata, "excel", confidence)
    
    def _process_json(self, input_data: Union[str, bytes, io.BytesIO]) -> ProcessedText:
        """Process JSON files."""
        try:
            if isinstance(input_data, bytes):
                json_text = input_data.decode('utf-8', errors='ignore')
            elif isinstance(input_data, io.BytesIO):
                json_text = input_data.read().decode('utf-8', errors='ignore')
            else:
                json_text = input_data
            
            # Parse and pretty-print JSON
            json_data = json.loads(json_text)
            text_content = "JSON Data:\n\n" + json.dumps(json_data, indent=2)
            
            metadata = {
                "format": "json",
                "size_bytes": len(json_text)
            }
            confidence = 1.0
            
        except Exception as e:
            text_content = f"Error processing JSON: {str(e)}\n\nRaw content:\n{input_data}"
            metadata = {"error": str(e), "format": "json"}
            confidence = 0.3
        
        return self._finalize_processing(text_content, metadata, "json", confidence)
    
    def _process_text(self, input_data: Union[str, bytes, io.BytesIO]) -> ProcessedText:
        """Process plain text."""
        try:
            if isinstance(input_data, bytes):
                text_content = input_data.decode('utf-8', errors='ignore')
            elif isinstance(input_data, io.BytesIO):
                text_content = input_data.read().decode('utf-8', errors='ignore')
            else:
                text_content = str(input_data)
            
            metadata = {
                "format": "text",
                "character_count": len(text_content),
                "line_count": len(text_content.splitlines())
            }
            confidence = 1.0
            
        except Exception as e:
            text_content = f"Error processing text: {str(e)}"
            metadata = {"error": str(e), "format": "text"}
            confidence = 0.0
        
        return self._finalize_processing(text_content, metadata, "text", confidence)
    
    def _finalize_processing(self, content: str, metadata: Dict[str, Any], 
                           format_type: str, confidence: float) -> ProcessedText:
        """Finalize processing by chunking and calculating tokens."""
        
        # Calculate total tokens
        total_tokens = self._count_tokens(content)
        
        # Create chunks if content is too large
        chunks = self._create_chunks(content, total_tokens)
        
        # Update metadata
        metadata.update({
            "total_tokens": total_tokens,
            "chunk_count": len(chunks),
            "processing_confidence": confidence
        })
        
        return ProcessedText(
            content=content,
            metadata=metadata,
            chunks=chunks,
            total_tokens=total_tokens,
            format_type=format_type,
            confidence=confidence
        )
    
    def _count_tokens(self, text: str) -> int:
        """Count tokens in text."""
        if self.tokenizer:
            try:
                return len(self.tokenizer.encode(text))
            except Exception:
                pass
        
        # Fallback: rough estimation
        return len(text.split()) * 1.3
    
    def _create_chunks(self, content: str, total_tokens: int) -> List[str]:
        """Create chunks from content if it exceeds max tokens."""
        
        if total_tokens <= self.max_tokens_per_chunk:
            return [content]
        
        # Split into sentences first
        sentences = self._split_into_sentences(content)
        chunks = []
        current_chunk = ""
        current_tokens = 0
        
        for sentence in sentences:
            sentence_tokens = self._count_tokens(sentence)
            
            # If adding this sentence would exceed limit, start new chunk
            if current_tokens + sentence_tokens > self.max_tokens_per_chunk and current_chunk:
                chunks.append(current_chunk.strip())
                current_chunk = sentence
                current_tokens = sentence_tokens
            else:
                current_chunk += " " + sentence
                current_tokens += sentence_tokens
        
        # Add final chunk
        if current_chunk.strip():
            chunks.append(current_chunk.strip())
        
        # If no chunks were created (e.g., single massive sentence), force split
        if not chunks:
            chunk_size = len(content) // ((total_tokens // self.max_tokens_per_chunk) + 1)
            chunks = [content[i:i+chunk_size] for i in range(0, len(content), chunk_size)]
        
        return chunks
    
    def _split_into_sentences(self, text: str) -> List[str]:
        """Split text into sentences."""
        # Simple sentence splitting (could be improved with NLTK)
        import re
        
        # Split on sentence endings
        sentences = re.split(r'[.!?]+\s+', text)
        
        # Split on newlines as well
        result = []
        for sentence in sentences:
            if '\n' in sentence:
                result.extend(sentence.split('\n'))
            else:
                result.append(sentence)
        
        return [s.strip() for s in result if s.strip()]
    
    def get_processing_stats(self, processed_text: ProcessedText) -> Dict[str, Any]:
        """Get comprehensive processing statistics."""
        return {
            "format_type": processed_text.format_type,
            "total_tokens": processed_text.total_tokens,
            "chunks": len(processed_text.chunks),
            "confidence": processed_text.confidence,
            "metadata": processed_text.metadata,
            "estimated_cost_tokens": processed_text.total_tokens * len(processed_text.chunks),
            "processing_complexity": self._assess_complexity(processed_text)
        }
    
    def _assess_complexity(self, processed_text: ProcessedText) -> str:
        """Assess text processing complexity."""
        if processed_text.confidence < 0.5:
            return "high_error_rate"
        elif len(processed_text.chunks) > 10:
            return "high_volume"
        elif processed_text.format_type in ['pdf', 'excel']:
            return "medium_structured"
        else:
            return "low_simple" 